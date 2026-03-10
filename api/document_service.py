"""
Serviço para processamento de documentos com chunking
"""

import os
import uuid
import logging
import numpy as np
import asyncio
from datetime import datetime
from typing import Dict, Any, List

logger = logging.getLogger(__name__)
from sqlalchemy.orm import Session

from embeddings_service import EmbeddingsService
from qdrant_service import QdrantService
from database import Document

CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "512"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "50"))
DATA_DIR = os.getenv("DATA_DIR", "/data")
STORAGE_DIR = os.getenv("STORAGE_DIR", os.path.join(DATA_DIR, "storage"))
# Documentos com menos tokens são tratados como 1 chunk único
MIN_TOKENS_FOR_CHUNKING = int(os.getenv("MIN_TOKENS_FOR_CHUNKING", "500"))
CHARS_PER_TOKEN = 4  # Aproximação: ~1 token = 4 caracteres
EMBEDDINGS_TIMEOUT_SECONDS = float(os.getenv("EMBEDDINGS_TIMEOUT_SECONDS", "300"))


class DocumentService:
    """Serviço para processar e armazenar documentos com chunking"""

    def __init__(self):
        os.makedirs(STORAGE_DIR, exist_ok=True)

    def _estimate_tokens(self, text: str) -> int:
        """Estima número de tokens por caracteres (~1 token = 4 chars)"""
        if not text:
            return 0
        return max(1, len(text) // CHARS_PER_TOKEN)

    def _split_into_chunks(self, text: str) -> List[str]:
        """
        Divide texto em chunks com overlap.
        Usa aproximação por caracteres (1 token ~ 4 chars).
        Tenta quebrar em limites de palavra quando possível.
        """
        if not text or not text.strip():
            return [""]
        chunks = []
        chunk_chars = CHUNK_SIZE * CHARS_PER_TOKEN
        overlap_chars = CHUNK_OVERLAP * CHARS_PER_TOKEN
        step = chunk_chars - overlap_chars
        start = 0
        while start < len(text):
            end = min(start + chunk_chars, len(text))
            chunk = text[start:end]
            if end < len(text):
                last_space = chunk.rfind(" ")
                if last_space > chunk_chars // 2:
                    chunk = chunk[: last_space + 1]
                    end = start + len(chunk)
            chunks.append(chunk.strip())
            if end >= len(text):
                break
            start = end - overlap_chars
        return [c for c in chunks if c]

    def _create_chunks(self, text_content: str) -> List[str]:
        """
        Decide se usa 1 chunk ou múltiplos baseado no tamanho.
        Retorna lista de strings (conteúdo de cada chunk).
        """
        est_tokens = self._estimate_tokens(text_content)
        if est_tokens < MIN_TOKENS_FOR_CHUNKING:
            return [text_content] if text_content.strip() else [""]
        return self._split_into_chunks(text_content)

    async def process_document(
        self,
        filename: str,
        content: bytes,
        db: Session,
        embeddings_service: EmbeddingsService,
        qdrant_service: QdrantService,
    ) -> Dict[str, Any]:
        """
        Processar documento completo:
        1. Salvar arquivo no filesystem local
        2. Extrair conteúdo de texto
        3. Gerar embedding
        4. Salvar metadados no PostgreSQL e vetores no Qdrant
        
        Args:
            filename: Nome do arquivo
            content: Conteúdo binário do arquivo
            db: Sessão do banco de dados
            embeddings_service: Serviço de embeddings
            
        Returns:
            Dicionário com informações do documento processado
        """
        # Gerar caminho único para o arquivo
        file_id = str(uuid.uuid4())
        file_ext = os.path.splitext(filename)[1]
        file_path = os.path.join(STORAGE_DIR, f"{file_id}{file_ext}")

        await asyncio.to_thread(self._write_file, file_path, content)
        
        # Extrair conteúdo de texto
        text_content = await asyncio.to_thread(self._extract_text, content, filename)

        # Dividir em chunks (1 chunk se pequeno, múltiplos se grande)
        chunk_texts = await asyncio.to_thread(self._create_chunks, text_content)
        chunk_texts = [c for c in chunk_texts if c and c.strip()]
        if not chunk_texts:
            chunk_texts = [text_content.strip() or "(conteúdo vazio)"]

        # Gerar embeddings em batch (mais eficiente).
        # ATENÇÃO: não usamos fallback de zeros — vetor inválido quebra a busca semântica.
        # Se o modelo falhar, o upload é rejeitado com erro claro para o cliente.
        try:
            embeddings = await asyncio.wait_for(
                asyncio.to_thread(embeddings_service.generate_embeddings_batch, chunk_texts),
                timeout=EMBEDDINGS_TIMEOUT_SECONDS,
            )
        except asyncio.TimeoutError as exc:
            logger.error(
                "Timeout ao gerar embeddings para '%s' (limite: %ss)",
                filename, EMBEDDINGS_TIMEOUT_SECONDS,
            )
            raise RuntimeError(
                f"Timeout ao gerar embeddings para '{filename}' "
                f"(limite: {EMBEDDINGS_TIMEOUT_SECONDS}s). "
                "O modelo pode estar carregando; tente novamente em alguns segundos."
            ) from exc
        except Exception as exc:
            logger.error(
                "Falha ao gerar embeddings para '%s': %s",
                filename, exc, exc_info=True,
            )
            raise RuntimeError(
                f"Falha ao gerar embeddings para '{filename}': {exc}"
            ) from exc

        # Validar que nenhum vetor retornou zerado
        norms = np.linalg.norm(embeddings, axis=1)
        zero_count = int(np.sum(norms == 0))
        if zero_count > 0:
            logger.error(
                "Embeddings inválidos (vetor zero) para '%s': %d/%d chunks",
                filename, zero_count, len(chunk_texts),
            )
            raise RuntimeError(
                f"{zero_count} chunks de '{filename}' retornaram vetor zero. "
                "Verifique o serviço de embeddings."
            )

        # Preparar metadados base
        metadata_base = {
            "file_size": len(content),
            "file_type": file_ext,
        }
        first_document_id = None
        now = datetime.utcnow()
        points = []

        for idx, (chunk_content, embedding) in enumerate(zip(chunk_texts, embeddings)):
            parent_id = first_document_id if idx > 0 else None
            doc = Document(
                filename=filename,
                file_path=file_path,
                content=chunk_content,
                embedding=None,
                extra_metadata=metadata_base,
                chunk_index=idx,
                parent_doc_id=parent_id,
                created_at=now,
                updated_at=now,
            )
            db.add(doc)
            db.flush()
            doc_id = doc.id
            if idx == 0:
                first_document_id = doc_id
            points.append(
                {
                    "id": doc_id,
                    "vector": embedding.tolist(),
                    "payload": {
                        "document_id": doc_id,
                        "filename": filename,
                        "chunk_index": idx,
                        "file_path": file_path,
                    },
                }
            )

        await qdrant_service.upsert_points(points)
        db.commit()

        return {
            "document_id": first_document_id,
            "filename": filename,
            "file_path": file_path,
            "content_length": len(text_content),
            "embedding_dimension": len(embeddings[0]) if len(embeddings) else 0,
            "chunk_count": len(chunk_texts),
        }

    def _write_file(self, path: str, content: bytes) -> None:
        with open(path, "wb") as f:
            f.write(content)
    
    def _extract_text(self, content: bytes, filename: str) -> str:
        """
        Extrair texto de diferentes formatos de arquivo
        
        Args:
            content: Conteúdo binário
            filename: Nome do arquivo (para determinar tipo)
            
        Returns:
            Texto extraído
        """
        file_ext = os.path.splitext(filename)[1].lower()
        
        if file_ext == '.txt':
            # Arquivo de texto simples
            try:
                return content.decode('utf-8')
            except UnicodeDecodeError:
                return content.decode('latin-1')
        
        elif file_ext == '.md':
            # Markdown
            try:
                return content.decode('utf-8')
            except UnicodeDecodeError:
                return content.decode('latin-1')
        
        elif file_ext == '.pdf':
            # PDF - requer biblioteca adicional
            try:
                import PyPDF2
                from io import BytesIO
                pdf_file = BytesIO(content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
                # Cleanup PyPDF2 artifacts (spaces and newlines between words)
                import re
                text = re.sub(r'\n[ \t]*\n?', ' ', text)
                text = re.sub(r' +', ' ', text)
                return text.strip()
            except ImportError:
                return "PDF requer PyPDF2. Instale com: pip install PyPDF2"
            except Exception as e:
                return f"Erro ao extrair PDF: {str(e)}"
        
        elif file_ext == '.docx':
            # DOCX - requer biblioteca adicional
            try:
                from docx import Document as DocxDocument
                from io import BytesIO
                docx_file = BytesIO(content)
                doc = DocxDocument(docx_file)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                return text
            except ImportError:
                return "DOCX requer python-docx. Instale com: pip install python-docx"
            except Exception as e:
                return f"Erro ao extrair DOCX: {str(e)}"
        
        else:
            # Tentar decodificar como texto genérico
            try:
                return content.decode('utf-8')
            except UnicodeDecodeError:
                return content.decode('latin-1', errors='ignore')

